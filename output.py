import numpy as np,files
import pandas as pd

def by_voting(paths):
    dataset=[pandas.read_csv(path_i)
                for path_i in paths]
    vote_dicts={}
    key_names=['Dataset','Clf']
    for data_i in dataset:
        for vote_j in data_i.voting.unique():
            data_ij=data_i[data_i['Voting']==vote_j]	
            vote_dicts[vote_j]=to_dict(data_ij,key_names)
    return vote_dicts

def to_dict(data,key_names):
    data = data.reset_index()  
    data_dict={}
    for index, row in data.iterrows():
        key_i="_".join([row[name_i] 
            for name_i in key_names])
        data_dict[key_i]=row
    return data_dict

def compare_output(pair,attr,vote_dicts):
    old,new=vote_dicts[pair[0]],vote_dicts[pair[1]]
    all_diff=[]
    for name_i in old.keys():
    	old_attr=old[name_i][attr]
    	new_attr=new[name_i][attr]
    	diff= new_attr-old_attr
    	print(f"{name_i},{diff/old_attr}")
    	all_diff.append( diff)
    print(np.median(all_diff))

def best_attr(in_path,attr='auc_mean'):
    exp_output = pandas.read_csv(in_path)
    show_attr=['Dataset','Clf','Voting']
    for name_i in exp_output.Dataset.unique():
        sub=exp_output[exp_output.Dataset==name_i]
        index=sub[attr].idxmax()
        max_value=sub[attr].max()
        row_i= exp_output.iloc[[index]]
        desc_i=[row_i[attr_i].values[0]  
                   for attr_i in show_attr]
        print(f"{desc_i},{max_value}")

def to_doc(in_path,out_path,group_size=3):
    from docx import Document
    exp_output = pandas.read_csv(in_path)
    cols= exp_output.columns
    n_groups=int(len(cols)/group_size)
    groups= [cols[(i*group_size):(i+1)*group_size]  
                for i in range(n_groups)]
    start,metrics=groups[0],groups[1:]
    
    document = Document()
    dataset_dict=by_dataset(exp_output)
 
    for name_i,lines_i in dataset_dict.items():
        by_metric=[[] for _  in metrics]
        for line_j in lines_i:
            for k,m_k in enumerate(metrics):
                unifed=list(line_j[start])+list(line_j[m_k])
                by_metric[k].append(unifed)
        for j,group_j in enumerate(by_metric):
            col_names_j=list(start)+list(metrics[j])
            table_i=document.add_table(rows=len(group_j)+1,
                  cols=len(col_names_j))
            fill_rows(table_i,0,col_names_j)
            for t,line_t in enumerate(group_j):
                fill_rows(table_i,t+1,line_t)
            document.add_paragraph()   
    document.save(out_path)

def by_dataset(exp_output):
    dataset_dict={ name_i:[]
       for name_i in exp_output.Dataset.unique()}
    for i, row_i in exp_output.iterrows():
        dataset_dict[row_i['Dataset']].append(row_i)
    return dataset_dict

def fill_rows(table_i,j,values):
    cells = table_i.rows[j].cells
    for i,value_i in enumerate(values):
        cells[i].text=str(value_i)    

def find_best(in_path,comp_path):
    base_df = pandas.read_csv(in_path)
    comp_df =as_dataframe(comp_path)
    best_result={}
    for i,row_i in comp_df.iterrows():
        dataset_i=row_i[0]
        row_i=pandas.to_numeric(row_i[1:])
        arg_i=row_i.argmax()
        clf_i=comp_df.columns[arg_i]
        best_result[dataset_i]=(clf_i,row_i[arg_i])
    s_cols=['Clf','Voting','Acc_mean']
    for dataset_i,(clf_i,value) in best_result.items():
        df_i=base_df[base_df.Dataset==dataset_i]
        if(len(df_i)>0):
            arg_max= df_i.Acc_mean.argmax()
            row_i=df_i.iloc[arg_max]
            acc_i=100*row_i['Acc_mean']
            base_i=",".join([str(row_i[col_j]) 
                        for col_j in s_cols])
            comp_i=f"{dataset_i},{clf_i},{value}"
            line_i=f"{comp_i},{base_i},{acc_i>value}"
            print(line_i)

def as_dataframe(output,cols=None):
    lines=[]
    for name_i,stats_i in output:
        lines.append( [name_i] + stats_i)
    return pd.DataFrame(lines,columns=cols)

if __name__ == "__main__":
    find_best("final/raw.csv","final/pruning.csv")